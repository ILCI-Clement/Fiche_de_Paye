"""Regression checks for attendance totals, snapshots and export transactions."""

from copy import deepcopy
from io import BytesIO
from pathlib import Path
from unittest import TestCase, main
from unittest.mock import patch
from zipfile import ZipFile

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from app_core import create_person, create_form, validate_form, ensure_workspace, find_duplicate_form, move_to_trash, restore_from_trash, import_legacy_forms
from attendance import monthly_attendance, internship_amounts
from export_service import build_export, build_individual_zip, prepare_exports, export_fingerprint


def sample_form(person_type='employee', month=9):
    person = create_person(person_type, last_name='Martin', first_name='Élodie', supervisor_name='William', start_date='2026-01-01', end_date='2026-12-31', hourly_rate=4.35, day_count=20.0, daily_hours=7.0, transport_cost=75.0, transport_rate=50.0)
    return create_form(2026, month, person)


class AttendanceTests(TestCase):
    def test_standard_month_and_custom_half_day(self):
        form = sample_form()
        self.assertEqual(sum(day['minutes'] for day in monthly_attendance(form)), 22 * 7 * 60)
        form['schedule']['Mardi']['afternoon_end'] = '18:30'
        form['exceptions']['2026-09-01'] = {'morning': {'type': 'paid_leave'}}
        day = monthly_attendance(form)[0]
        self.assertEqual(day['minutes'], 330)
        self.assertEqual(day['afternoon']['label'], '13:00–18:30')

    def test_holiday_override_and_weekend(self):
        form = sample_form(month=5)
        self.assertEqual(monthly_attendance(form)[0]['minutes'], 0)
        form['exceptions']['2026-05-01'] = {'morning': {'type': 'work'}}
        self.assertEqual(monthly_attendance(form)[0]['minutes'], 180)
        form['schedule']['Samedi']['active'] = True
        self.assertEqual(monthly_attendance(form)[1]['minutes'], 420)

    def test_snapshot_never_follows_person(self):
        person = create_person('employee', last_name='Martin')
        form = create_form(2026, 9, person)
        person['default_schedule']['Lundi']['morning_start'] = '10:00'
        self.assertEqual(form['schedule']['Lundi']['morning_start'], '09:00')

    def test_validation_rejects_reversed_time_outside_month_and_other(self):
        form = sample_form()
        self.assertFalse(validate_form(form))
        form['schedule']['Lundi']['morning_start'] = '14:00'
        form['exceptions']['2026-10-01'] = {'morning': {'type': 'other', 'hours': -1, 'label': ''}}
        self.assertGreaterEqual(len(validate_form(form)), 3)

    def test_blank_and_person_duplicate(self):
        first = sample_form()
        second = create_form(2026, 9, person_type='employee', last_name='  martin ', first_name='Élodie')
        self.assertIs(find_duplicate_form([first], second), first)

    def test_restore_preserves_archive_and_blocks_duplicate(self):
        form = sample_form()
        form.update(state='archived', archived=True)
        workspace = ensure_workspace({'forms': [form]})
        move_to_trash(workspace, form['id'])
        restore_from_trash(workspace, form['id'])
        self.assertEqual(form['state'], 'archived')
        move_to_trash(workspace, form['id'])
        workspace['forms'].append(sample_form())
        with self.assertRaises(ValueError):
            restore_from_trash(workspace, form['id'])
        self.assertEqual(len(workspace['trash']), 1)

    def test_legacy_import_is_repeatable_without_data_loss(self):
        source = {'type': 'Salarié', 'nom': 'Martin Élodie', 'responsable': 'William', 'ddc': '2026-01-01', 'fdc': '2026-12-31', 'vacances': [{'date': '2026-09-01', 'matin': True, 'aprem': False}]}
        workspace = ensure_workspace({'mois': 9, 'annee': 2026, 'employes_data': [source]})
        original = deepcopy(source)
        self.assertEqual(import_legacy_forms(workspace), 1)
        self.assertEqual(import_legacy_forms(workspace), 0)
        self.assertEqual(workspace['employes_data'][0], original)
        self.assertEqual(workspace['forms'][0]['exceptions']['2026-09-01']['morning']['type'], 'paid_leave')

    def test_intern_amounts_retain_manual_basis(self):
        form = sample_form('intern')
        self.assertEqual(str(internship_amounts(form)['total']), '646.50')
        form['exceptions']['2026-09-01'] = {'morning': {'type': 'absence'}}
        self.assertEqual(str(internship_amounts(form)['total']), '646.50')


class ExportTests(TestCase):
    def test_employee_excel_uses_template_and_month_total(self):
        workbook = load_workbook(BytesIO(build_export('excel', [sample_form()])))
        sheet = workbook.active
        self.assertEqual(sheet['AA4'].value, 'UNIVERSAL EDUCATION GROUP')
        self.assertAlmostEqual(sheet['AF25'].value.total_seconds() / 3600, 154)
        self.assertEqual(sheet['I30'].value, 0)
        self.assertIsNone(sheet['W30'].value)

    def test_word_contains_all_dates_and_no_employee_pay(self):
        document = Document(BytesIO(build_export('word', [sample_form()])))
        text = '\n'.join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        self.assertIn('30/09', text)
        self.assertIn('154:00', text)
        self.assertNotIn('Taux horaire', text)

    def test_original_intern_template_is_rendered(self):
        document = Document(BytesIO(build_export('word', [sample_form('intern')])) )
        text = '\n'.join([paragraph.text for paragraph in document.paragraphs] + [cell.text for table in document.tables for row in table.rows for cell in row.cells])
        self.assertIn('646.50', text)
        self.assertNotIn('{{', text)
        self.assertIn('Universal Education Group', text)
        self.assertIn('30/09', text)
        self.assertIn('154:00', text)

    def test_intern_excel_contains_monthly_attendance_table(self):
        workbook = load_workbook(BytesIO(build_export('excel', [sample_form('intern')])) )
        sheet = workbook.active
        self.assertEqual([sheet.cell(22, column).value for column in range(1, 5)], ['Date', 'Matin', 'Après-midi', 'Heures'])
        self.assertEqual(sheet['A23'].value.strftime('%d/%m'), '01/09')
        self.assertEqual(sheet['A52'].value.strftime('%d/%m'), '30/09')
        self.assertAlmostEqual(sheet['D53'].value.total_seconds() / 3600, 154)

    def test_pdf_full_month_and_two_people(self):
        pdf = PdfReader(BytesIO(build_export('pdf', [sample_form(), sample_form('intern')])) )
        self.assertEqual(len(pdf.pages), 2)
        self.assertIn('30/09', pdf.pages[0].extract_text())
        self.assertIn('646.50', pdf.pages[1].extract_text())
        self.assertIn('30/09', pdf.pages[1].extract_text())

    def test_pdf_keeps_long_notes_and_last_day(self):
        form = sample_form()
        for day in range(1, 31):
            form['exceptions'][f'2026-09-{day:02d}'] = {'morning': {'type': 'other', 'label': 'Formation interne détaillée. ' * 10 + f'FIN-{day}', 'hours': 3}}
        pdf = PdfReader(BytesIO(build_export('pdf', [form])))
        text = '\n'.join(page.extract_text() for page in pdf.pages)
        self.assertIn('FIN-30', text)

    def test_zip_same_names_have_distinct_paths(self):
        archive = ZipFile(BytesIO(build_individual_zip([sample_form(), sample_form()], ['pdf'])))
        self.assertEqual(len(set(archive.namelist())), 2)
        self.assertTrue(all('/MARTIN_Élodie_09-2026_fiche_presence.pdf' in name for name in archive.namelist()))

    def test_individual_download_names_include_person_and_period(self):
        form = sample_form()
        package = prepare_exports([form], ['word', 'pdf', 'excel'])
        self.assertEqual(
            [file['name'] for file in package['files']],
            [
                'MARTIN_Élodie_09-2026_fiche_presence.docx',
                'MARTIN_Élodie_09-2026_fiche_presence.pdf',
                'MARTIN_Élodie_09-2026_fiche_presence.xlsx',
            ],
        )
        separate = prepare_exports([form], ['pdf'], separate=True)
        self.assertEqual(separate['files'][0]['name'], 'MARTIN_Élodie_09-2026_fiche_presence_exports.zip')

    def test_same_period_batch_download_name_starts_with_period(self):
        package = prepare_exports([sample_form(), sample_form()], ['pdf'])
        self.assertEqual(package['files'][0]['name'], '09-2026_fiches.pdf')

    def test_export_failure_does_not_mutate_states(self):
        form = sample_form()
        original = deepcopy(form)
        with patch('export_service.build_export', side_effect=[b'valid', ValueError('broken')]):
            with self.assertRaises(ValueError):
                prepare_exports([form], ['excel', 'pdf'])
        self.assertEqual(form, original)

    def test_input_changes_invalidate_prepared_bytes(self):
        form = sample_form()
        original = export_fingerprint([form])
        form['state'] = 'exported'
        self.assertEqual(export_fingerprint([form]), original)
        form['last_name'] = 'Dupont'
        self.assertNotEqual(export_fingerprint([form]), original)

    def test_excel_user_text_cannot_become_formula(self):
        form = sample_form()
        form['last_name'] = '=1+1'
        sheet = load_workbook(BytesIO(build_export('excel', [form]))).active
        self.assertEqual(sheet['AA2'].data_type, 's')


if __name__ == '__main__':
    main()
