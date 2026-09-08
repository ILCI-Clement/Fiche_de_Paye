"""Export builders for attendance and internship allowance forms."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
from collections import Counter
from copy import deepcopy
from html import escape
import hashlib
import json

from docx import Document
from docxtpl import DocxTemplate
from num2words import num2words
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import load_workbook
from openpyxl.cell.cell import MergedCell
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from app_core import validate_form
from attendance import monthly_attendance, format_minutes, internship_amounts

TEMPLATE_PATH = Path(__file__).resolve().parent / 'Fiche_Exemple.xlsx'


EXPORT_MIME_TYPES = {
    "excel": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "word": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}
EXPORT_EXTENSIONS = {"excel": "xlsx", "word": "docx", "pdf": "pdf"}
EXPORT_LABELS = {"excel": "Excel", "word": "Word", "pdf": "PDF"}


def safe_name(value: str) -> str:
    """Create a filename-safe French label."""
    name = "_".join("".join(char if char.isalnum() or char in "-_" else " " for char in value).split()) or "sans_nom"
    if name.upper() in {'CON', 'PRN', 'AUX', 'NUL', *(f'COM{i}' for i in range(1, 10)), *(f'LPT{i}' for i in range(1, 10))}:
        name = '_' + name
    return name[:100]


def display_name(form: dict) -> str:
    """Return the employee or intern name for an export."""
    return " ".join(
        part for part in (form.get("last_name", "").upper(), form.get("first_name", "")) if part
    ) or "Sans nom"


def form_title(form: dict) -> str:
    """Return the French title for a form."""
    person_label = "Fiche de présence" if form.get("person_type") == "employee" else "Fiche d’indemnité de stage"
    return f"{person_label} — {form['month']:02d}/{form['year']}"


def individual_export_stem(form: dict) -> str:
    """Prefix an individual export with the person's name and month-year."""
    person_prefix = safe_name(display_name(form))[:70]
    document_label = 'fiche_presence' if form.get('person_type') == 'employee' else 'fiche_indemnite_stage'
    return f"{person_prefix}_{form['month']:02d}-{form['year']}_{document_label}"


def individual_export_filename(form: dict, format_name: str) -> str:
    """Return the final filename for one person and one format."""
    return f"{individual_export_stem(form)}.{EXPORT_EXTENSIONS[format_name]}"


def combined_export_filename(forms: list[dict], format_name: str) -> str:
    """Name one-person exports precisely and prefix same-period batches."""
    if len(forms) == 1:
        return individual_export_filename(forms[0], format_name)
    periods = {(form['year'], form['month']) for form in forms}
    prefix = f'{forms[0]["month"]:02d}-{forms[0]["year"]}' if len(periods) == 1 else 'multi-periodes'
    return f'{prefix}_fiches.{EXPORT_EXTENSIONS[format_name]}'


def individual_package_filename(forms: list[dict]) -> str:
    """Name the downloadable ZIP consistently with its individual files."""
    if len(forms) == 1:
        return f'{individual_export_stem(forms[0])}_exports.zip'
    periods = {(form['year'], form['month']) for form in forms}
    prefix = f'{forms[0]["month"]:02d}-{forms[0]["year"]}' if len(periods) == 1 else 'multi-periodes'
    return f'{prefix}_exports_individuels.zip'


def form_rows(form: dict) -> list[tuple[str, str]]:
    """Build common export rows from a monthly form snapshot."""
    snapshot = form.get("person_snapshot", {})
    rows = [
        ("Personne", display_name(form)),
        ("Type", "Salarié" if form.get("person_type") == "employee" else "Stagiaire"),
        ("Période", f"{form['month']:02d}/{form['year']}"),
        ("Responsable", snapshot.get("supervisor_name", "")),
        ("E-mail du responsable", snapshot.get("supervisor_email", "")),
        ("Début du contrat / stage", str(snapshot.get('start_date') or '')),
        ("Fin du contrat / stage", 'CDI' if snapshot.get('permanent_contract') and form.get('person_type') == 'employee' else str(snapshot.get('end_date') or '')),
    ]
    rows.append(('Heures du mois', format_minutes(sum(day['minutes'] for day in monthly_attendance(form)))))
    if form['person_type'] == 'intern':
        amounts = internship_amounts(form)
        rows += [('Taux horaire (€)', str(snapshot.get('hourly_rate', 0))),
                 ('Jours indemnisés', str(snapshot.get('day_count', 0))),
                 ('Heures indemnisées / jour', str(snapshot.get('daily_hours', 0))),
                 ('Indemnité de stage (€)', str(amounts['stage'])),
                 ('Transport', snapshot.get('transport', '')),
                 ('Facture transport (€)', str(snapshot.get('transport_cost', 0))),
                 ('Remboursement (%)', str(snapshot.get('transport_rate', 0))),
                 ('Indemnité transport (€)', str(amounts['transport'])),
                 ('Total indemnité (€)', str(amounts['total']))]
    return rows


def build_intern_word(form):
    """Render the unchanged internship template using its original placeholders."""
    details = form['person_snapshot']
    amounts = internship_amounts(form)
    months = ('Janvier', 'Février', 'Mars', 'Avril', 'Mai', 'Juin', 'Juillet', 'Août', 'Septembre', 'Octobre', 'Novembre', 'Décembre')
    template = DocxTemplate(TEMPLATE_PATH.with_name('template_stagiaire.docx'))
    context = {'nom': form['last_name'].upper(), 'prenom': form['first_name'], 'mois': months[form['month']-1],
               'annee': form['year'], 'dds': details['start_date'], 'fds': details['end_date'],
               'taux_horaire': f"{details.get('hourly_rate', 0):.2f}", 'nb_jours': details.get('day_count', 0),
               'nb_heures_jour': details.get('daily_hours', 0), 'total_stage': str(amounts['stage']),
               'transport': details.get('transport', ''), 'facture_mensuelle': f"{details.get('transport_cost', 0):.2f}",
               'taux': f"{details.get('transport_rate', 0)}%", 'total_transport': str(amounts['transport']),
               'total': str(amounts['total']), 'total_lettres': num2words(amounts['total'], lang='fr', to='currency').capitalize()}
    template.render(context, autoescape=True)
    output = BytesIO()
    template.save(output)
    return output.getvalue()


def apply_word_table_grid(table) -> None:
    """Draw borders even when the source template has no built-in table style."""
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '808080')
        borders.append(element)
    table._tbl.tblPr.append(borders)


def set_excel_text(cell, value):
    """Keep user text literal, even when it starts with a formula marker."""
    cell.value = str(value)
    cell.data_type = 's'


def write_excel_sheet(sheet, form: dict) -> None:
    """Populate the existing six-week template without changing its source file."""
    details = form.get('person_snapshot', {})
    if form['person_type'] == 'intern':
        sheet.merge_cells('A1:D1')
        sheet['A1'] = form_title(form)
        sheet['A1'].font = Font(size=14, bold=True)
        sheet['A1'].alignment = Alignment(horizontal='center')
        for index, (label, value) in enumerate(form_rows(form), 3):
            sheet.cell(index, 1, label).font = Font(bold=True)
            set_excel_text(sheet.cell(index, 2), value)
            sheet.merge_cells(start_row=index, end_row=index, start_column=2, end_column=4)

        header_row = 22
        headers = ('Date', 'Matin', 'Après-midi', 'Heures')
        border = Border(
            left=Side(style='thin', color='808080'),
            right=Side(style='thin', color='808080'),
            top=Side(style='thin', color='808080'),
            bottom=Side(style='thin', color='808080'),
        )
        for column, label in enumerate(headers, 1):
            cell = sheet.cell(header_row, column, label)
            cell.font = Font(bold=True)
            cell.fill = PatternFill('solid', fgColor='D9EAF7')
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border
        attendance = monthly_attendance(form)
        for row, day in enumerate(attendance, header_row + 1):
            values = (day['date'], day['morning']['label'], day['afternoon']['label'], day['minutes'] / 1440)
            for column, value in enumerate(values, 1):
                cell = sheet.cell(row, column, value)
                cell.border = border
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            sheet.cell(row, 1).number_format = 'dd/mm/yyyy'
            sheet.cell(row, 4).number_format = '[h]:mm'
        total_row = header_row + len(attendance) + 1
        sheet.merge_cells(start_row=total_row, end_row=total_row, start_column=1, end_column=3)
        sheet.cell(total_row, 1, 'Total des heures').font = Font(bold=True)
        sheet.cell(total_row, 1).alignment = Alignment(horizontal='right')
        sheet.cell(total_row, 4, sum(day['minutes'] for day in attendance) / 1440)
        sheet.cell(total_row, 4).number_format = '[h]:mm'
        for column in range(1, 5):
            sheet.cell(total_row, column).border = border

        sheet.column_dimensions['A'].width = 22
        sheet.column_dimensions['B'].width = 28
        sheet.column_dimensions['C'].width = 28
        sheet.column_dimensions['D'].width = 14
        sheet.freeze_panes = f'A{header_row + 1}'
        sheet.page_setup.orientation = 'landscape'
        sheet.page_setup.paperSize = sheet.PAPERSIZE_A4
        sheet.sheet_properties.pageSetUpPr.fitToPage = True
        sheet.page_setup.fitToWidth = sheet.page_setup.fitToHeight = 1
        sheet.print_area = f'A1:D{total_row}'
        return
    for address, value in {'AA2': display_name(form), 'AA6': details.get('supervisor_name', ''),
                           'M2': f"{form['month']:02d}", 'M5': form['year'],
                           'AC30': str(details.get('start_date') or ''),
                           'AC35': 'CDI' if details.get('permanent_contract') and form['person_type'] == 'employee' else str(details.get('end_date') or ''),
                           'AF14': '', 'AF17': ''}.items():
        set_excel_text(sheet[address], value)
    for address in ('U30', 'W30', 'U35', 'W35'):
        sheet[address] = None
    for row in sheet.iter_rows(min_row=11, max_row=24, min_col=2, max_col=31):
        for cell in row:
            if not isinstance(cell, MergedCell):
                cell.value = None
    days = monthly_attendance(form)
    week_totals = [0] * 6
    counts = Counter()
    for day in days:
        week = (days[0]['date'].weekday() + day['date'].day - 1) // 7
        column = 2 + 5 * week
        row = 11 + day['date'].weekday() * 2
        sheet.cell(row, column, day['date'].day)
        for offset, half in enumerate(('morning', 'afternoon')):
            item = day[half]
            if item['type'] == 'work':
                sheet.cell(row + offset, column + 1, item['start'])
                sheet.cell(row + offset, column + 2, 'à')
                sheet.cell(row + offset, column + 3, item['end'])
            else:
                sheet.merge_cells(start_row=row + offset, end_row=row + offset, start_column=column + 1, end_column=column + 3)
                cell = sheet.cell(row + offset, column + 1)
                set_excel_text(cell, item['label'])
                cell.alignment = Alignment(wrap_text=True, horizontal='center', vertical='center')
                sheet.row_dimensions[row + offset].height = max(25, sheet.row_dimensions[row + offset].height or 0)
            sheet.cell(row + offset, column + 4, item['minutes'] / 1440).number_format = '[h]:mm'
            counts[item['type']] += 0.5
        week_totals[week] += day['minutes']
    for index, total in enumerate(week_totals):
        sheet.cell(25, 2 + 5 * index, total / 1440).number_format = '[h]:mm'
    sheet['AF25'] = sum(week_totals) / 1440
    sheet['AF25'].number_format = '[h]:mm'
    for address, status in [('I30', 'paid_leave'), ('I27', 'absence'), ('S27', 'sick_leave')]:
        sheet[address] = counts[status]
    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    sheet.page_setup.orientation = 'landscape'
    sheet.page_setup.paperSize = sheet.PAPERSIZE_A4
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 1
    sheet.print_area = 'A1:AG37'


def build_excel(forms: list[dict]) -> bytes:
    """Build an Excel workbook with one sheet for every form."""
    workbook = load_workbook(TEMPLATE_PATH)
    template = workbook.active
    for form in forms:
        sheet = workbook.copy_worksheet(template) if form['person_type'] == 'employee' else workbook.create_sheet()
        sheet.title = safe_name(display_name(form))[:25] + f'_{len(workbook.worksheets)-1}'
        write_excel_sheet(sheet, form)
    workbook.remove(template)
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def append_word_attendance_table(document: Document, form: dict, *, page_break=False) -> None:
    """Append the same monthly attendance table for employees and interns."""
    if page_break:
        document.add_page_break()
    document.add_heading("Présences du mois", level=2)
    exception_table = document.add_table(rows=1, cols=4)
    try:
        exception_table.style = "Table Grid"
    except KeyError:
        apply_word_table_grid(exception_table)
    for index, label in enumerate(("Date", "Matin", "Après-midi", 'Heures')):
        exception_table.rows[0].cells[index].text = label
    for day in monthly_attendance(form):
        cells = exception_table.add_row().cells
        cells[0].text = day['date'].strftime('%d/%m')
        cells[1].text = day['morning']['label']
        cells[2].text = day['afternoon']['label']
        cells[3].text = format_minutes(day['minutes'])
    total_cells = exception_table.add_row().cells
    total_cells[0].merge(total_cells[1]).merge(total_cells[2])
    total_cells[0].text = 'Total des heures'
    total_cells[3].text = format_minutes(sum(day['minutes'] for day in monthly_attendance(form)))


def write_word_form(document: Document, form: dict) -> None:
    """Append one French form to a Word document."""
    document.add_heading(form_title(form), level=1)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in form_rows(form):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    append_word_attendance_table(document, form)


def build_intern_word_with_attendance(form: dict) -> bytes:
    """Preserve the internship allowance template and append its attendance table."""
    document = Document(BytesIO(build_intern_word(form)))
    append_word_attendance_table(document, form, page_break=True)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_word(forms: list[dict]) -> bytes:
    """Build a Word document with one form per page."""
    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = section.right_margin = Cm(1.4)
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(8)
    normal_style.paragraph_format.space_after = Pt(0)
    for index, form in enumerate(forms):
        if index:
            document.add_page_break()
        if form['person_type'] == 'intern':
            if len(forms) == 1:
                return build_intern_word_with_attendance(form)
            part = Document(BytesIO(build_intern_word_with_attendance(form)))
            for element in list(part.element.body):
                if not element.tag.endswith('}sectPr'):
                    document.element.body.insert(-1, deepcopy(element))
        else:
            write_word_form(document, form)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_pdf(forms: list[dict]) -> bytes:
    """Render every day and wrap long descriptions; never discard overflow."""
    output = BytesIO()
    styles = getSampleStyleSheet()
    styles['Normal'].fontSize = 8
    styles['Normal'].leading = 10
    story = []
    def paragraph(text):
        return Paragraph(escape(str(text)), styles['Normal'])
    for index, form in enumerate(forms):
        if index:
            story.append(PageBreak())
        story.append(Paragraph(escape(form_title(form)), styles['Heading2']))
        story.extend(paragraph(f'{label} : {value}') for label, value in form_rows(form))
        story.append(Paragraph('Présences du mois', styles['Heading3']))
        rows = [[paragraph(value) for value in ('Date', 'Matin', 'Après-midi', 'Heures')]]
        rows += [[paragraph(day['date'].strftime('%d/%m')), paragraph(day['morning']['label']), paragraph(day['afternoon']['label']), paragraph(format_minutes(day['minutes']))] for day in monthly_attendance(form)]
        rows.append([paragraph('Total des heures'), '', '', paragraph(format_minutes(sum(day['minutes'] for day in monthly_attendance(form))))])
        table = Table(rows, colWidths=[1.5*cm, 7*cm, 7*cm, 1.5*cm], repeatRows=1)
        table.setStyle(TableStyle([('GRID', (0,0), (-1,-1), .4, colors.grey), ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#eeeeee')), ('VALIGN', (0,0), (-1,-1), 'TOP'), ('SPAN', (0,-1), (2,-1)), ('ALIGN', (0,-1), (2,-1), 'RIGHT'), ('TOPPADDING', (0,0), (-1,-1), 2), ('BOTTOMPADDING', (0,0), (-1,-1), 2)]))
        story.append(table)
    SimpleDocTemplate(output, pagesize=A4, topMargin=1.2*cm, bottomMargin=1.2*cm, leftMargin=1.4*cm, rightMargin=1.4*cm).build(story)
    return output.getvalue()


def build_export(format_name: str, forms: list[dict]) -> bytes:
    """Build one export file in the selected format."""
    if not forms:
        raise ValueError('Sélectionnez au moins une fiche.')
    for form in forms:
        errors = validate_form(form)
        if errors:
            raise ValueError(f"{display_name(form)} : " + ' '.join(errors))
    if format_name == "excel":
        return build_excel(forms)
    if format_name == "word":
        return build_word(forms)
    if format_name == "pdf":
        return build_pdf(forms)
    raise ValueError("Unsupported export format")


def build_individual_zip(forms: list[dict], formats: list[str]) -> bytes:
    """Build a ZIP with one directory per person and one file per format."""
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as archive:
        for index, form in enumerate(forms, 1):
            person_folder = f'{safe_name(display_name(form))}_{index:03d}'
            for format_name in formats:
                filename = individual_export_filename(form, format_name)
                archive.writestr(f"{person_folder}/{filename}", build_export(format_name, [form]))
    return output.getvalue()


def export_fingerprint(forms):
    """Ignore workflow state when checking whether prepared bytes still match inputs."""
    fields = ('id', 'first_name', 'last_name', 'person_type', 'person_snapshot', 'year', 'month', 'schedule', 'exceptions')
    payload = [{field: form.get(field) for field in fields} for form in forms]
    return hashlib.sha256(json.dumps(payload, sort_keys=True, default=str).encode()).hexdigest()


def prepare_exports(forms, formats, separate=False):
    """Build atomically: callers replace downloads only after every format succeeds."""
    if not forms or not formats:
        raise ValueError('Sélectionnez des fiches et au moins un format.')
    if separate:
        files = [{'name': individual_package_filename(forms), 'mime': 'application/zip', 'data': build_individual_zip(forms, formats), 'formats': list(formats)}]
    else:
        files = [{'name': combined_export_filename(forms, format_name), 'mime': EXPORT_MIME_TYPES[format_name], 'data': build_export(format_name, forms), 'formats': [format_name]} for format_name in formats]
    return {'ids': [form['id'] for form in forms], 'fingerprint': export_fingerprint(forms), 'files': files}
